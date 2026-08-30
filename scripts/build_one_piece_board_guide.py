from __future__ import annotations

import json
import re
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ARTIFACT_ROOT = ROOT / "artifacts" / "one_piece_board_complete_guide_v1"
DEFAULT_REPORT = ARTIFACT_ROOT / "qa" / "complete_run" / "full_campaign_2p2cpu_report.json"
DEFAULT_ANNOTATED = ARTIFACT_ROOT / "qa" / "complete_run" / "screenshots" / "annotated"
DEFAULT_OUTPUT = ARTIFACT_ROOT / "航海王大富翁_一周目二周目全Boss超詳細攻略_V1.docx"
ITEM_CATALOG = ROOT / "docs" / "ITEM_CATALOG.md"
ARSENAL_REPORT = ARTIFACT_ROOT / "qa" / "bullet_arsenal" / "bullet_arsenal_full_compatibility_result.json"

NAVY = "071923"
NAVY_2 = "0C2B3A"
TEAL = "1F8A8A"
GOLD = "D7A928"
PALE_GOLD = "FFF5D6"
PALE_TEAL = "E7F4F2"
PALE_BLUE = "EDF4F8"
RED = "A62C2B"
GRAY = "53636B"
LIGHT = "F7F9FA"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge_name, edge_data in edges.items():
        edge = tc_borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            tc_borders.append(edge)
        for key, value in edge_data.items():
            edge.set(qn(f"w:{key}"), str(value))


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:cantSplit")
    tr_pr.append(node)


def set_east_asia(run, font_name="Microsoft JhengHei"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_run(run, size=None, bold=None, color=None, font="Microsoft JhengHei"):
    set_east_asia(run, font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def setup_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft JhengHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(9.2)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color, before, after in (
        ("Title", 31, NAVY, 0, 10),
        ("Subtitle", 13, TEAL, 0, 12),
        ("Heading 1", 20, NAVY, 14, 8),
        ("Heading 2", 14, TEAL, 10, 5),
        ("Heading 3", 11, RED, 7, 3),
    ):
        style = doc.styles[style_name]
        style.font.name = "Microsoft JhengHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_run(header.add_run("ONE PIECE BOARD｜一周目＋二周目全攻略"), 7.5, True, TEAL)
        footer = section.footer.paragraphs[0]
        set_run(footer.add_run("實測攻略 V1　｜　"), 7.5, False, GRAY)
        add_page_number(footer)


def add_rule(doc, color=GOLD):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "14")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def add_callout(doc, title: str, body: str, tone="teal"):
    fill = {"teal": PALE_TEAL, "gold": PALE_GOLD, "blue": PALE_BLUE, "red": "FBEAEA"}.get(tone, PALE_TEAL)
    accent = {"teal": TEAL, "gold": GOLD, "blue": "4C7A92", "red": RED}.get(tone, TEAL)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 105, 150, 105, 150)
    set_cell_border(cell, left={"val": "single", "sz": "22", "color": accent})
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run(title), 10, True, accent)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    set_run(p2.add_run(body), 8.8, False, NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullets(doc, lines, level=0):
    for text in lines:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.paragraph_format.space_after = Pt(2)
        set_run(p.add_run(str(text)), 9)


def add_numbered(doc, lines):
    for text in lines:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        set_run(p.add_run(str(text)), 9)


def add_picture(doc, path: Path, caption: str, width=6.65):
    if not path.exists():
        add_callout(doc, "截圖缺少", f"找不到：{path.name}", "red")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    shape = p.add_run().add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", caption)
    shape._inline.docPr.set("title", path.stem)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6)
    set_run(cap.add_run(caption), 7.8, False, GRAY)


def add_simple_table(doc, headers, rows, widths=None, font_size=8.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, label in enumerate(headers):
        cell = header.cells[i]
        set_cell_shading(cell, NAVY_2)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            cell.width = Inches(widths[i])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(str(label)), font_size, True, WHITE)
    for row_index, values in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        for i, value in enumerate(values):
            cell = row.cells[i]
            set_cell_margins(cell, 55, 65, 55, 65)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index % 2:
                set_cell_shading(cell, LIGHT)
            if widths:
                cell.width = Inches(widths[i])
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run(p.add_run(str(value)), font_size, i == 0, NAVY)
    return table


def parse_markdown_items(path: Path):
    text = path.read_text(encoding="utf-8")
    sections = {}
    current = None
    for line in text.splitlines():
        if line.startswith("## ") and "（" in line:
            current = re.sub(r"（\d+）", "", line[3:]).strip()
            sections[current] = []
            continue
        if not current or not line.startswith("|"):
            continue
        cells = [cell.strip().replace("`", "") for cell in line.strip().strip("|").split("|")]
        if len(cells) != 6 or cells[0] in {"名稱", "---"} or set(cells[0]) == {"-"}:
            continue
        if cells[0].startswith("---"):
            continue
        sections[current].append({
            "name": cells[0], "id": cells[1], "rarity": cells[2],
            "effect": cells[3], "source": cells[4], "price": cells[5],
        })
    return sections


def safe_text(value, fallback="—"):
    text = str(value or "").strip()
    return text or fallback


def find_shot(report, annotated_dir: Path, includes):
    if isinstance(includes, str):
        includes = [includes]
    for shot in report.get("screenshots", []):
        label = str(shot.get("label", ""))
        if all(token in label for token in includes):
            return annotated_dir / shot["file"]
    return Path("__missing__")


def boss_screenshot(report, annotated_dir: Path, boss_name: str):
    aliases = {
        "薩卡／七星劍": ["薩卡"],
        "綠牛／荒牧": ["綠牛"],
        "洛克斯・D・吉貝克": ["洛克斯終戰"],
        "Tot Musica": ["Tot", "Musica"],
    }
    tokens = aliases.get(boss_name, [boss_name.replace("・", "")])
    for shot in report.get("screenshots", []):
        label = str(shot.get("label", "")).replace("・", "")
        if all(token.replace("・", "") in label for token in tokens):
            return annotated_dir / shot["file"]
    return Path("__missing__")


def add_chapter(doc, number: int, title: str, deck: str = ""):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run(f"CHAPTER {number:02d}"), 9, True, GOLD)
    doc.add_heading(title, level=1)
    if deck:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        set_run(p.add_run(deck), 10.5, False, TEAL)
    add_rule(doc)


def add_cover(doc, report, annotated_dir):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, NAVY)
    set_cell_margins(cell, 330, 260, 260, 260)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("ONE PIECE BOARD"), 13, True, GOLD)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(8)
    set_run(p2.add_run("航海王大富翁\n一周目＋二周目全 Boss\n超詳細攻略書"), 28, True, WHITE)
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p3.add_run("2 真人＋2 CPU 全流程實測版"), 12, True, "7FE0D8")
    add_picture(doc, find_shot(report, annotated_dir, "二周目十三孤島地圖"), "二周目世界地圖：13 座無風帶孤島在每局的 Boss 配置順序可能不同。", 6.55)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("版本 V1　｜　2026-08-29　｜　依現行正式程式與實測畫面編製"), 8.5, False, GRAY)
    doc.add_page_break()


def add_front_matter(doc, report):
    doc.add_heading("這本攻略怎麼讀", level=1)
    add_callout(
        doc,
        "實測範圍",
        "建立真實 Socket.IO 房間，配置 2 名真人與 CPU1、CPU2，走過四人輪替、四皇、伊姆、一周目結局、13 位二周目 Boss、洛基、約克解碼與洛克斯。戰鬥進場、同步、血統因子、掉落與返回地圖都走正式流程；僅把 Boss HP 歸零這一步用確定勝利加速。",
        "gold",
    )
    add_bullets(doc, [
        "想快速通關：先看第 1～4 章，再直接翻到遇到的 Boss。",
        "遇到畫面像卡住：先看第 10 章；很多時候是在等待開場被動、獎勵揭示或下一位玩家接手。",
        "二周目 13 位 Boss 的島嶼順序每局可能不同，所以用 Boss 名稱找章節，不要只看孤島編號。",
        "所有新挑戰都以滿 HP 開場；只有同一場續戰才保留先前 HP、PP 與 Boss 進度。",
    ])

    doc.add_heading("章節導航", level=2)
    rows = [
        ("01", "四人開局與快速準備", "2 真人＋2 CPU、輪替、房主操作"),
        ("02", "航海與地圖", "路線、事件、商店、醫院與高價道具"),
        ("03", "戰鬥核心", "四指令、屬性、速度、骰子、暴擊、續戰"),
        ("04", "一周目", "四皇、四拓本、伊姆、最終之島"),
        ("05", "二周目 13 Boss", "每位 Boss 一頁以上的機制與反制"),
        ("06", "二周目終局", "洛基、約克解碼、洛克斯、研究所"),
        ("07", "多人玩法", "組隊、共鬥、推進城救援、PK 切磋"),
        ("08", "208 件道具總表", "取得方式與商店價格"),
        ("09", "配裝與武器庫", "4,371 種組合掃描後的實用搭配"),
        ("10", "卡住排查與 QA", "刷新、等待狀態與完整測試矩陣"),
    ]
    add_simple_table(doc, ["章", "主題", "用途"], rows, [0.5, 2.0, 4.3], 8.5)

    room = report.get("room", {})
    add_callout(
        doc,
        "本次完整房間證據",
        f"房號 {safe_text(room.get('roomCode'))}；席位 {len(room.get('players', []))}；CPU {sum(1 for p in room.get('players', []) if p.get('isCPU'))}。完整報告記錄 {len(report.get('screenshots', []))} 張關鍵畫面。",
        "blue",
    )


def add_chapter_one(doc, report, annotated_dir):
    add_chapter(doc, 1, "四人開局與快速準備", "先讓四個席位、輪到誰與操作權都明確，後面的同步才會穩。")
    add_picture(doc, find_shot(report, annotated_dir, "2真人2CPU等待室"), "實測等待室：兩位真人、CPU1、CPU2 同時在房。")
    doc.add_heading("開房順序", level=2)
    add_numbered(doc, [
        "房主建立航海房，把房號交給第二名真人。",
        "第二名真人加入後按「準備」；房主再加入 CPU1、CPU2。",
        "房主確認四個席位與名稱正確後開始遊戲。CPU 不需要真人代按準備。",
        "進入選角後，每位玩家完成自己的操作；正式主地圖會依既有玩家順序交棒。",
    ])
    add_callout(doc, "誰能按？", "只讓目前回合玩家操作。觀戰玩家看得到同一場戰鬥，但不能趁同步中替別人選招。CPU 回合由房主端代為推進。", "teal")
    add_picture(doc, find_shot(report, annotated_dir, "四人主地圖"), "四人主地圖：左上看目前玩家，船圖周圍的呼吸光代表可操作目標。")
    doc.add_heading("四人輪替實測", level=2)
    ring = report.get("turnRing", [])
    rows = []
    for entry in ring:
        rows.append((entry.get("before", {}).get("name", ""), "→", entry.get("after", {}).get("name", ""), entry.get("round", "")))
    add_simple_table(doc, ["本回合", "", "下一位", "回合數"], rows, [2.2, 0.35, 2.2, 1.0], 8.5)
    add_bullets(doc, [
        "CPU 倍速只縮短等待，不改骰點、事件結果或掉落判定。",
        "多人房如果輪到別人，你可以觀看；等輪到自己才提取血統因子、看個人結算或處理延後共鬥獎勵。",
        "看到按鈕不能點時，先確認左上目前玩家，以及船圖是否有可操作呼吸光。",
    ])


def add_chapter_two(doc):
    add_chapter(doc, 2, "航海、島嶼與資源規劃", "前期的真正難點是經濟與續航，不是把每一場遭遇都打到底。")
    doc.add_heading("每回合優先檢查", level=2)
    add_simple_table(doc, ["順序", "檢查", "決策"], [
        ("1", "隊伍 HP／PP", "不滿就評估醫院、回復道具或避免高階戰鬥。"),
        ("2", "本回合路線", "先看島嶼服務、未知格與可用航海道具。"),
        ("3", "背包", "指定步數券、單雙數骰子很強，但價格高，留給關鍵節點。"),
        ("4", "船隻被動", "船隻工具在商店販售，越強的改造越晚買得起。"),
        ("5", "目前玩家", "多人時由當前玩家抽事件與決定路線。"),
    ], [0.55, 1.8, 4.35], 8.5)
    doc.add_heading("事件與服務", level=2)
    add_bullets(doc, [
        "醫院：全員狀態圓滿時不會再抽到伊娃而白白浪費回合。",
        "商店：不鎖玩家等級；以價格形成前中後期門檻。單數骰子、雙數骰子與指定步數券屬高價戰術品。",
        "酒館／招募：留意船上最多六名角色；同一角色 ID 不能同時重複登船。",
        "船隻工具：16 件都可在一般道具商店購買；先買穩定續航，再追求爆發航行。",
        "海格戰：一般戰可撤退；Boss 戰與特殊戰通常要依機制處理。",
    ])
    add_callout(doc, "前期購物建議", "低價解異常藥、帶骨肉與固定續航船具優先；4,800 B 的單／雙數骰子、3,800 B 的指定步數券適合保留到 Boss、關鍵航路或二周目機制。", "gold")


def add_chapter_three(doc):
    add_chapter(doc, 3, "戰鬥核心：四指令、骰子與暴擊", "所有一般戰與多數 Boss 仍維持四顆主指令，不額外塞第五顆按鈕。")
    doc.add_heading("四顆主指令", level=2)
    add_simple_table(doc, ["指令", "用途", "常見誤判"], [
        ("攻擊", "選擇目前角色已解鎖招式；PP 會正式消耗。", "多段連擊仍是直接攻擊，不會因為叫『連擊』就失去同步資格。"),
        ("夥伴", "主動換人或瀕死替補。", "強制替補、交棒旗與普通換人有不同時機。"),
        ("道具", "治療、解除異常、控制骰子或指定對象。", "卡塔庫栗正面出招後仍可用道具，不必硬攻。"),
        ("逃跑", "一般戰撤退；特殊 Boss 依規則限制。", "成功逃離敵島後，下次航行要沿原路返回。"),
    ], [0.75, 2.35, 3.6], 8.1)
    doc.add_heading("屬性克制", level=2)
    add_callout(doc, "克制循環", "力 → 技、技 → 速、速 → 力。九尾幻面可把攻擊屬性轉成克制目前敵人的屬性；如果原本已經克制，就不重播發動畫面。", "teal")
    doc.add_heading("速度、先攻與行動優先度", level=2)
    add_bullets(doc, [
        "一般戰先比較招式優先度，再比較有效速度。",
        "Tot Musica 同時出戰兩名角色時，先攻速度取當前兩位中較快者；雙世界行動優先度取較高者。",
        "疾風圍巾目前只套用一次，基準實測約 ×1.252；不會再重複疊成舊版約 ×1.565。",
    ])
    doc.add_heading("暴擊", level=2)
    add_bullets(doc, [
        "每位玩家角色都有依人物形象設定的基礎暴擊率，並隨升等、進化提高；敵方本體不會暴擊。",
        "敵人／Boss 透過血統因子成為我方角色後，會取得自己的玩家版暴擊率；洛克斯 Lv.1 基礎為 30%。",
        "傷害在正式命中結算時逐段判定暴擊；多段攻擊可只有其中幾段暴擊。",
        "技能按鈕只在說明裡寫「10%暴擊」「20%暴擊」，不顯示額外角標；狀態列不常駐暴擊率與暴傷。",
        "實測上限：暴擊率 50%，暴擊傷害倍率最高 ×2。",
    ])
    add_callout(doc, "新挑戰與續戰", "新開一場挑戰會把隊伍補成滿 HP；重整後回到同一場 pending battle 則保留當時 HP、PP、狀態、Boss HP 與機制計數。", "blue")


FIRST_BOSS_GUIDE = {
    "黑鬍子": {
        "reward": "戴彭的九尾幻面",
        "plan": "先穩定命中與解控制，再用短回合爆發；不要讓黑暗壓制拖成資源戰。",
        "counter": "後續再次挑戰時，太陽海賊團徽章能對黑鬍子維持破暗穴與削弱震波的反制價值。",
    },
    "紅髮香克斯": {
        "reward": "格里芬之劍",
        "plan": "觀察力／技／速相位；只有克制目前相位的攻擊能造成傷害。九尾幻面可大幅降低換招壓力。",
        "counter": "保留不同屬性招式或可換屬性的角色，避免整隊被單一相位封死。",
    },
    "大媽": {
        "reward": "太陽海賊團的徽章",
        "plan": "保留生命與控制抗性，不要把全隊壓到同時容易被靈魂拷問的血線。",
        "counter": "無畏之心對大媽有 80% 靈魂拷問抵抗，並能以最大 HP 固定傷害縮短戰鬥。",
    },
    "凱多": {
        "reward": "無畏之心",
        "plan": "用速屬性、骰到 6，或破防／破盾／破霸體招式破壞龍鱗；先破鱗再交爆發。",
        "counter": "格里芬之劍可撕開防禦，對凱多額外破壞龍鱗。",
    },
    "伊姆": {
        "reward": "最終之島航路與一周目結局",
        "plan": "準備全隊回復、解除控制與能處理強化的招式；不要在深淵壓制後硬用低命中高成本招。",
        "counter": "完成四皇並備齊四份路標歷史本文後，從背包啟動最終之島；打贏伊姆後才進結局。",
    },
}


def add_first_playthrough(doc, report, annotated_dir):
    add_chapter(doc, 4, "一周目：四皇、伊姆與最終之島", "跑完最終之島結局才算完成一周目；四皇全倒只是最後航路的前置。")
    yonko = report.get("firstPlaythrough", {}).get("yonko", [])
    for index, boss in enumerate(yonko, 1):
        name = boss.get("bossName", "四皇")
        doc.add_heading(f"4-{index}　{name}", level=2)
        add_picture(doc, boss_screenshot(report, annotated_dir, name), f"一周目實測：{name} 正式戰鬥進場。", 6.45)
        guide = FIRST_BOSS_GUIDE.get(name, {})
        add_simple_table(doc, ["項目", "內容"], [
            ("實測 HP", f"{boss.get('maxHp', '—')}（正式平衡器會依現行遊戲狀態調整）"),
            ("屬性／定位", f"{safe_text(boss.get('attribute'))}／{safe_text(boss.get('role'))}"),
            ("作戰重點", guide.get("plan", "先處理 Boss 專屬機制，再交爆發。")),
            ("反制", guide.get("counter", "保持全隊 HP 與 PP，避免把所有資源押在單一角色。")),
            ("首次關鍵獎勵", guide.get("reward", "路標歷史本文")),
        ], [1.25, 5.55], 8.3)

    doc.add_page_break()
    doc.add_heading("4-5　四份拓本啟動最終之島", level=2)
    add_picture(doc, find_shot(report, annotated_dir, "四拓本啟動最終之島"), "四份路標歷史本文集齊後，從背包正式啟動。")
    add_picture(doc, find_shot(report, annotated_dir, "最終之島航路現身"), "啟動後航路與最終節點才出現在主地圖。")
    imu = report.get("firstPlaythrough", {}).get("imu", {})
    doc.add_heading("4-6　伊姆最終戰", level=2)
    add_picture(doc, boss_screenshot(report, annotated_dir, "伊姆"), "伊姆戰：確認 HP、控制狀態與可用回復，再開始正式出招。")
    guide = FIRST_BOSS_GUIDE["伊姆"]
    add_simple_table(doc, ["項目", "內容"], [
        ("實測 HP", safe_text(imu.get("hp"))),
        ("屬性", safe_text(imu.get("attribute"))),
        ("招式", "、".join(move.get("name", "") for move in imu.get("moves", []))),
        ("作戰重點", guide["plan"]),
    ], [1.1, 5.7], 8.3)
    doc.add_heading("4-7　一周目結局與二周目解鎖", level=2)
    add_picture(doc, find_shot(report, annotated_dir, "一周目最終之島結局"), "結局劇情播放完成後，二周目無風帶孤島與研究線才正式解鎖。")
    add_callout(doc, "判定點", "只打贏伊姆但沒有完成最終之島結局，不算正式進入二周目。跳過劇情可以，但必須讓結局流程完成解鎖與交棒。", "red")


CANONICAL_BOSS_KEYS = [
    "postgame_shiki", "postgame_gild_tesoro", "postgame_zephyr", "postgame_tot_musica",
    "postgame_douglas_bullet", "postgame_saga", "postgame_vinsmoke_judge",
    "postgame_rob_lucci_awakened", "postgame_king", "postgame_charlotte_katakuri",
    "postgame_patrick_redfield", "postgame_oars", "postgame_aramaki",
]


def boss_extra_tips(key):
    tips = {
        "postgame_shiki": ["優先拆岩獅島，否則直攻本體只保留極低傷害。", "三島全毀後再集中高威力招式；連擊可把剩餘段數轉到下一島。"],
        "postgame_gild_tesoro": ["每名船員分開記錄三段金流；換到後排會清空。", "第三次命中若同時致死，走一般擊倒，不再多一次金流強制換人。"],
        "postgame_zephyr": ["倒數只有 4；前期至少安排一次解除，不要只貪傷害。", "點左側炸藥岩進行解除，不新增第五顆按鈕。累積 3 點才停止倒數。"],
        "postgame_tot_musica": ["兩邊都必須選直接攻擊，第一骰同奇偶才完整同步。", "同屬性、同第一骰、同攻擊類型三項可疊到 ×2.2；一個攻擊、一個連擊仍是直接攻擊。"],
        "postgame_douglas_bullet": ["可空裝入場，避免讓巴雷特吸到高價攜帶物。", "武器庫被吸收只形成空外殼；裡面兩件留在原持有人身上但暫不生效，破壞該孔才恢復。"],
        "postgame_saga": ["血祭值只會增加，不會被治療倒扣。", "護盾與解除流血比純補血更有效；75 點融合後要快速收尾。"],
        "postgame_vinsmoke_judge": ["開場三名複製兵，每段命中各擋一段。", "三段以上連擊能一次清完人牆；別讓伽治完成三次行動補兵。"],
        "postgame_rob_lucci_awakened": ["每回合先看六式：剃／鐵塊／紙繪／月步／指槍／嵐腳各強化不同面向。", "抽到第六式只是六王銃預警；真正輪到路基才施放必中、無視 50% 防禦的重擊。"],
        "postgame_king": ["點燃背火時只承受 10% 傷害，不要浪費大招。", "觀察每次行動後的狀態切換；熄火回合防護消失，但 King 輸出與速度更高。"],
        "postgame_charlotte_katakuri": ["先看預知骰，再選正面出招或防禦。", "正面出招仍可換人、用道具、逃跑；只有實際招式的第一骰會和預知骰比較。"],
        "postgame_patrick_redfield": ["第一骰幾點，就吸每名存活船員最大 HP 幾％。", "降命中、提高閃避或控制他；攻擊落空／完全擋下就不吸血。"],
        "postgame_oars": ["歐斯 HP 55,555，但速度與命中極低；可慢慢打，也可玩鹽袋。", "共享鹽袋達 15 包後用巨型鹽彈直接淨化。換人／道具／逃跑沒有攻擊骰，下注會退回。"],
        "postgame_aramaki": ["先清森林再打到 0 HP，否則第一次倒下會依森林片數復活。", "高骰清 2 片；火焰可額外清林，但防火林會擋住加成。"],
    }
    return tips.get(key, ["先讀機制計數器再選招。", "保留一名高血量替補，避免機制失敗時全隊同時倒下。"])


def add_postgame_bosses(doc, report, annotated_dir):
    add_chapter(doc, 5, "二周目：13 位無風帶 Boss", "島嶼編號是每局配置；攻略順序依 Boss 名稱整理。每次全新挑戰固定滿血。")
    add_picture(doc, find_shot(report, annotated_dir, "二周目十三孤島地圖"), "13 座無風帶孤島已出現在主地圖；每座對應一位 Boss 與一張約克線索。")
    bosses = {entry.get("bossKey"): entry for entry in report.get("secondPlaythrough", {}).get("bosses", [])}
    for index, key in enumerate(CANONICAL_BOSS_KEYS, 1):
        boss = bosses.get(key)
        if not boss:
            continue
        doc.add_page_break()
        name = boss.get("bossName", key)
        doc.add_heading(f"5-{index:02d}　{name}", level=2)
        subtitle = f"HP {boss.get('maxHp', '—')}　｜　屬性 {safe_text(boss.get('attribute'))}　｜　{safe_text(boss.get('role'))}"
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(subtitle), 9.5, True, TEAL)
        add_picture(doc, boss_screenshot(report, annotated_dir, name), f"實測戰鬥畫面：{name}。黃色編號標示 HP／機制／四指令。", 6.35)
        mechanic = boss.get("mechanic") or {}
        add_callout(doc, safe_text(mechanic.get("title"), "Boss 機制"), safe_text(mechanic.get("rule"), boss.get("passiveText")), "gold")
        doc.add_heading("攻略重點", level=3)
        add_bullets(doc, boss_extra_tips(key))
        if mechanic.get("counter"):
            add_callout(doc, "遊戲內反制提示", mechanic["counter"], "teal")
        moves = boss.get("moves") or []
        if moves:
            rows = []
            for move in moves:
                rows.append((move.get("name", ""), move.get("type", ""), move.get("power", 0) or "—", move.get("description", "") or "依戰鬥狀態套用"))
            add_simple_table(doc, ["招式", "類型", "威力", "效果摘要"], rows, [1.55, 0.75, 0.55, 3.85], 7.7)
        add_callout(doc, "勝利後", "取得對應約克線索；專屬攜帶物獨立進行 10% 掉落判定。島上培養艙重建後，必須先離島再重新登島才能再次挑戰。", "blue")


def add_postgame_finale(doc, report, annotated_dir):
    add_chapter(doc, 6, "洛基、約克解碼與洛克斯終戰", "13 張線索是蛋頭島的鑰匙；洛基是獨立試煉，不占 13 Boss 線索。")
    second = report.get("secondPlaythrough", {})
    loki = second.get("loki") or {}
    doc.add_heading("6-1　洛基王子試煉", level=2)
    add_picture(doc, boss_screenshot(report, annotated_dir, "洛基"), "洛基屬於可重複挑戰的艾爾巴夫試煉。")
    add_bullets(doc, [
        f"實測 HP {loki.get('hp', '—')}，屬性 {safe_text(loki.get('attribute'))}。",
        "勝利可進行血統因子抽取，並獨立判定鐵雷 Ragnir；不會增加約克 13 張線索。",
        "CPU 完成基本整備後能自動選擇挑戰，不會卡在島嶼選擇視窗。",
    ])
    doc.add_heading("6-2　約克解碼", level=2)
    add_picture(doc, find_shot(report, annotated_dir, "約克解碼後蛋頭島現身"), "13／13 線索與解碼器條件達成後，蛋頭島與洛克斯終戰節點現身。")
    york = second.get("york") or {}
    add_simple_table(doc, ["檢查", "實測結果"], [
        ("線索種類", f"{york.get('clueEntries', 13)}／13"),
        ("可啟動", "是" if york.get("canActivate") else "否"),
        ("蛋頭島解鎖", "是" if york.get("eggheadUnlocked") else "否"),
        ("解碼器階級", york.get("decoderTier", "—")),
    ], [2.0, 4.8], 8.5)
    doc.add_heading("6-3　洛克斯・D・吉貝克", level=2)
    rocks = second.get("rocks") or {}
    add_picture(doc, boss_screenshot(report, annotated_dir, "洛克斯・D・吉貝克"), "二周目終局：洛克斯戰。先確認研究所與隊伍已整備。")
    add_bullets(doc, [
        f"實測屬性 {safe_text(rocks.get('attribute'))}；正式角色檔的玩家版洛克斯 Lv.1 基礎暴擊率 30%。",
        "招式包含日蝕、霸王色纏繞、世界之王的野心、神之谷崩裂。",
        "首次勝利獲得完美血統因子核心 ×1、研究點數 25，並解鎖神之谷霸王框。",
        "名刀『日蝕』依個人約克解碼器階級決定掉落率；可以重複取得，但不能交易。",
    ])
    add_callout(doc, "研究點數", "二周目 SSS 級正式勝利每場 25 點；研究等級門檻為 0／40／120／300。打贏洛克斯後仍要收下道具揭示，才會交棒給下一位玩家。", "gold")


def add_multiplayer(doc):
    add_chapter(doc, 7, "多人：組隊、共鬥、推進城與 PK", "多人不是共享一個按鈕；操作權、延後結算與主線順序都有清楚邊界。")
    doc.add_heading("同層組隊", level=2)
    add_bullets(doc, [
        "兩名玩家同時位於同一樓層時，輪到該樓層的玩家可發出邀請；對方同意後才組隊成功。",
        "事件與路線由目前回合玩家抽取、決定；隱藏囚犯則每位玩家輪到自己時個別抽，不被隊長獨占。",
        "遇到戰鬥時，同層每名玩家各行動一次，套用既有共鬥交棒。",
        "共鬥打完後，血統因子提取與個人結算仍要等到自己的地圖回合。觀戰者不能搶先領。",
    ])
    doc.add_heading("推進城待救援", level=2)
    add_simple_table(doc, ["狀況", "結果"], [
        ("某玩家全隊瀕死", "進入待救援 2 個個人回合。"),
        ("仍有同隊玩家存活", "戰鬥可繼續；倒下者等待救援。"),
        ("所有小隊全滅", "直接判定逃獄失敗。"),
        ("救援成功", "救出對方並延續逃獄。"),
        ("救援失敗", "救援者被押往海軍本部。"),
    ], [2.0, 4.8], 8.5)
    doc.add_heading("Tot Musica 的多人分隊", level=2)
    add_bullets(doc, [
        "新加入共鬥的玩家仍被分進現實世界或歌世界；每個世界可有多位玩家依序交棒。",
        "任一世界的整支小隊全滅，就因無法維持雙世界同步而直接戰敗。",
        "先攻速度取當前兩位出戰角色中較快者；行動優先度取兩個行動中較高者。",
        "兩邊都選直接攻擊才有同步判定；普通攻擊和連擊可以組合。",
    ])
    doc.add_heading("二周目 PK 切磋", level=2)
    add_bullets(doc, [
        "二周目同格玩家除了交易，也能發起切磋。雙方先在準備室看到對方船上六名角色，但看不到攜帶物。",
        "雙方秘密選三名上場；三名全倒才分勝負。PK 以滿 HP、滿 PP 開始，不影響主線角色 HP、PP、攜帶物與存檔。",
        "原順序 A→B→C→D 不被打破：A 與 C 每人交一輪指令後，地圖換 B；再到 C 時繼續下一輪，之後換 D。所有玩家都看得到戰鬥。",
        "非參戰者回合不建立全域戰鬥狀態；PK 戰況暫存在 activeSpar 快照，輪到參戰者時才恢復。",
    ])


def add_item_catalog(doc, sections):
    add_chapter(doc, 8, "208 件道具總表：取得方式與價格", "表內以現行正式資料為準；不可購買品會明確標示其掉落、事件或首次通關來源。")
    counts = [(name, len(rows)) for name, rows in sections.items()]
    add_simple_table(doc, ["類型", "件數"], counts + [("合計", sum(count for _, count in counts))], [4.6, 1.2], 9)
    add_callout(doc, "商店規則", "商店不鎖玩家等級，強力道具以價格形成門檻。16 件船隻道具全部可販售；Boss、任務、事件與掉落限定品則標為不可購買。", "gold")
    for category, items in sections.items():
        doc.add_page_break()
        doc.add_heading(f"8｜{category}（{len(items)}）", level=2)
        rows = []
        for item in items:
            name = f"{item['name']}\n[{item['id']}]"
            source_price = f"取得：{item['source']}\n價格：{item['price']}"
            rows.append((name, item["rarity"], item["effect"], source_price))
        add_simple_table(doc, ["名稱／ID", "稀有", "效果", "取得方式／商店價格"], rows, [1.45, 0.48, 2.25, 2.62], 7.0)


def load_arsenal_rankings():
    if not ARSENAL_REPORT.exists():
        return {}
    return json.loads(ARSENAL_REPORT.read_text(encoding="utf-8")).get("rankings", {})


def add_builds(doc, rankings):
    add_chapter(doc, 9, "角色攜帶物與巴雷特武器庫配裝", "武器庫可裝兩件不同攜帶物；本次把 95 件攜帶物的所有合法雙組合全掃過。")
    add_callout(doc, "武器庫限制", "只能裝兩件不同攜帶物；同一種不可重複。戰鬥 HUD 不把三件效果全攤開，點開武器庫再查看內裝與狀態。卸下武器庫時，內裝物一併回背包。", "teal")
    add_bullets(doc, [
        "巴雷特吸收武器庫時只取得空外殼，不展開裡面兩件，也不取得它們的效果。",
        "內裝兩件留在原持有人身上但暫時失效；破壞武器庫孔位後才恢復生效。",
        "香吉士同時裝傑爾馬66戰鬥服與九尾幻面時，九尾幻面負責屬性轉換；只有第二階新世界香吉士會由戰鬥服變身隱形黑。兩個演出依戰鬥時機排隊，不會同時塞爆畫面。",
    ])
    categories = [("general", "綜合"), ("offense", "輸出"), ("survival", "生存"), ("speed", "速度")]
    for key, label in categories:
        entries = rankings.get(key, [])[:8]
        if not entries:
            continue
        doc.add_heading(f"{label}向前 8 組", level=2)
        rows = []
        for rank, entry in enumerate(entries, 1):
            names = "＋".join(entry.get("names", []))
            scores = entry.get("scores", {})
            rows.append((rank, names, f"綜合 {scores.get('general', '—')}", f"輸出 {scores.get('offense', '—')}／生存 {scores.get('survival', '—')}／速度 {scores.get('speed', '—')}"))
        add_simple_table(doc, ["名次", "組合", "總評", "分項"], rows, [0.45, 2.75, 1.0, 2.6], 7.8)
    doc.add_heading("三套容易理解的實戰方向", level=2)
    add_simple_table(doc, ["方向", "建議組合", "理由"], [
        ("爆發", "櫻十・木枯＋Battle Smasher", "物理、速度與蓄熱爆發兼具；Smasher 反噬不能致死，落空仍消耗爆發。"),
        ("穩定生存", "傑爾馬66戰鬥服＋森森生命種子", "一次致命傷防護搭配低血回復／種子復活，容錯高。"),
        ("速度", "疾風圍巾＋覺醒黑焰羽衣", "速度乘區最高的一組之一；圍巾會鎖定首次攻擊招式，要先選好主力技。"),
        ("吸血護盾", "赤色伯爵的傘劍＋高傷武器", "整招結算後依實際傷害吸一次，滿血溢出轉最大 HP 15% 護盾。"),
    ], [1.0, 2.25, 3.55], 8.1)


def add_troubleshooting(doc, report):
    add_chapter(doc, 10, "卡住排查、重整規則與實測結論", "先判斷是在等待動畫、等待本人、等待揭示，還是真的沒有狀態前進。")
    doc.add_heading("看起來像卡住的正常等待", level=2)
    add_simple_table(doc, ["畫面", "先做什麼", "原因"], [
        ("戰前對話／被動畫面", "等演出完成，不要連點招式。", "開場對話、卡塔庫栗見聞色、六式、九尾幻面、隱形黑等會依序播放。"),
        ("Boss 已倒但還在戰鬥", "看是否跳出血統因子提取；選提取或放棄，再按返回地圖。", "研究所啟用後，戰後多一層個人決定。"),
        ("戰鬥關閉但回合沒換", "等道具卡出現「點擊繼續」，逐張收下。", "約克線索、Boss 專武、完美核心等重要道具會阻止偷偷交棒。"),
        ("共鬥其他人看不到結算", "等輪到那名玩家自己的地圖回合。", "多人共鬥獎勵與提取採各自回合處理。"),
        ("CPU 停在島嶼選擇", "先等一次 CPU 決策；若長時間不動再重整房主。", "醫院、商店與洛基都有 CPU 專用判斷。"),
    ], [1.45, 2.25, 3.1], 8.1)
    doc.add_heading("什麼時候可以重整", level=2)
    add_bullets(doc, [
        "戰鬥中重整：會從多人完整快照回到同一場戰鬥，保留 Boss、HP、PP、狀態、行動佇列與機制。",
        "輪到別人時重整：仍是觀戰者，不能因此取得操作權。",
        "若顯示『等待房主狀態』：先等房主在線；沒有快照時不會用本機新局覆蓋房間。",
        "不要連續多台同時重整。先讓一台恢復並看到正確目前玩家，再處理下一台。",
    ])
    doc.add_heading("本次專項測試矩陣", level=2)
    tests = [
        ("2 真人＋2 CPU 完整房", "四人席位、真人準備、CPU1／CPU2、四人輪替", "通過"),
        ("一周目", "四皇、四拓本、伊姆、結局、二周目解鎖", "通過"),
        ("二周目 13 Boss", "逐場進場、機制、滿血、結算、線索、培養艙重建", "通過"),
        ("終局", "洛基、13／13 約克、蛋頭島、洛克斯獎勵", "通過"),
        ("刷新恢復", "主回合、CPU 回合、戰鬥進場、觀戰端重整", "通過"),
        ("Tot Musica", "雙世界、同奇偶／不同奇偶、替補、全滅、血統提取", "通過"),
        ("共鬥延後結算", "只有輪到本人才能提取與收結算", "通過"),
        ("二周目 PK", "秘密選 3、滿狀態、主線隔離、A-B-C-D 回合順序", "通過"),
        ("暴擊", "51 名正式角色、66 種敵方轉我方、技能、道具、UI", "45／45 通過"),
        ("巴雷特武器庫", "95 件攜帶物、4,371 種不同雙組合", "全部通過"),
        ("桌機／平板", "戰鬥按鈕、角色卡、傷害與機制面板無水平溢出", "通過"),
    ]
    add_simple_table(doc, ["測試", "範圍", "結果"], tests, [1.3, 4.6, 0.9], 7.9)
    add_callout(doc, "本次修正範圍", "沒有改動正式戰鬥規則；修正的是 QA 腳本的同步等待、等待室可見判定、開場被動後才判定可操作，以及血統因子結算只續行一次，避免測試器自己製造假卡住。", "blue")
    final_sync = report.get("finalSync") or {}
    add_callout(
        doc,
        "最終同步確認",
        f"連線：{'正常' if final_sync.get('connected') else '異常'}；玩家 {final_sync.get('players', '—')}；CPU {final_sync.get('cpuPlayers', '—')}；二周目 {'已解鎖' if final_sync.get('postgameUnlocked') else '未解鎖'}；蛋頭島 {'已解鎖' if final_sync.get('eggheadUnlocked') else '未解鎖'}；殘留戰鬥 {'有' if final_sync.get('battleActive') else '無'}。",
        "gold",
    )


def main():
    report_path = Path(sys.argv[1]).resolve() if len(sys.argv) > 1 else DEFAULT_REPORT
    annotated_dir = Path(sys.argv[2]).resolve() if len(sys.argv) > 2 else DEFAULT_ANNOTATED
    output_path = Path(sys.argv[3]).resolve() if len(sys.argv) > 3 else DEFAULT_OUTPUT
    report = json.loads(report_path.read_text(encoding="utf-8"))
    if len(report.get("screenshots", [])) < 29:
        raise SystemExit(f"QA report is incomplete: {len(report.get('screenshots', []))} screenshots")
    if not report.get("ok"):
        raise SystemExit("QA report is not passing; refusing to label the guide as verified")
    sections = parse_markdown_items(ITEM_CATALOG)
    if sum(len(items) for items in sections.values()) != 208:
        raise SystemExit("Item catalog parse count is not 208")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_document(doc)
    add_cover(doc, report, annotated_dir)
    add_front_matter(doc, report)
    add_chapter_one(doc, report, annotated_dir)
    add_chapter_two(doc)
    add_chapter_three(doc)
    add_first_playthrough(doc, report, annotated_dir)
    add_postgame_bosses(doc, report, annotated_dir)
    add_postgame_finale(doc, report, annotated_dir)
    add_multiplayer(doc)
    add_item_catalog(doc, sections)
    add_builds(doc, load_arsenal_rankings())
    add_troubleshooting(doc, report)
    doc.core_properties.title = "航海王大富翁：一周目＋二周目全 Boss 超詳細攻略書"
    doc.core_properties.subject = "2 真人＋2 CPU 全流程實測與 208 件道具總表"
    doc.core_properties.author = "Codex × 王曜瑋"
    doc.core_properties.keywords = "One Piece Board, 一周目, 二周目, Boss, 攻略, 多人, CPU"
    doc.save(output_path)
    print(json.dumps({"output": str(output_path), "bosses": 20, "items": 208, "screenshots": len(report.get("screenshots", []))}, ensure_ascii=False))


if __name__ == "__main__":
    main()
